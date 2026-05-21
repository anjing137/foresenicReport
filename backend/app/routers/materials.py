"""
材料管理路由 - 支持按类型上传、按医院分组
"""
import os
import re
import shutil
import uuid
import json
from io import BytesIO
from difflib import SequenceMatcher
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, UploadFile, File, Form, Query, HTTPException, Body
from sqlalchemy.orm import Session
from typing import Optional, List

from app.database import get_db
from app.models.case import (
    HospitalNameAlias,
    HospitalRecord,
    ImagingReport,
    Material,
    MaterialGroup,
    MaterialType,
    MedicalEvent,
    OcrStatus,
    Report,
    UnifiedFact,
)
from app.schemas.case import (
    MaterialResponse, MaterialGroupResponse, MaterialGroupCreate, MaterialGroupUpdate
)
from app.config import settings
from app.utils.material_classifier import (
    CLASSIFIER_VERSION,
    classify_material_subtype,
    classify_pdf_page,
    classify_pdf_page_with_llm,
    get_material_subtype_label,
)
from app.utils.ocr import extract_text_from_result, run_ocr
from app.utils.source_material import material_original_page_number, material_sequence_key

router = APIRouter(prefix="/api/materials", tags=["材料管理"])

UPLOAD_DIR = str(settings.UPLOAD_DIR)
os.makedirs(UPLOAD_DIR, exist_ok=True)


def _clean_filename(filename: Optional[str], fallback: str = "upload") -> str:
    """保留展示用原始文件名，同时去掉路径片段和空值。"""
    cleaned = os.path.basename(filename or "").replace("\x00", "").strip()
    return cleaned or fallback


def _storage_filename(original_filename: str) -> str:
    """生成安全、唯一的落盘文件名，避免重名覆盖和特殊路径。"""
    stem, ext = os.path.splitext(original_filename)
    safe_stem = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", stem).strip("._")
    safe_stem = safe_stem[:80] or "file"
    safe_ext = re.sub(r"[^A-Za-z0-9.]", "", ext)[:16]
    return f"{safe_stem}_{uuid.uuid4().hex[:8]}{safe_ext.lower()}"


def _normalize_extracted_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _escape_markdown_cell(text: str) -> str:
    return (text or "").replace("\n", " ").replace("|", "\\|").strip()


def _extract_docx_markdown(path: str) -> str:
    """轻量解析 docx，保留段落和表格文本，供分类与后续提取复用。"""
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise HTTPException(500, f"缺少 python-docx 依赖，无法解析 Word 文档: {exc}")

    try:
        doc = DocxDocument(path)
    except Exception as exc:
        raise HTTPException(400, f"Word 文档解析失败，请确认文件为有效 .docx: {exc}")

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading" in style_name:
            level_match = re.search(r"(\d+)", style_name)
            level = min(int(level_match.group(1)), 6) if level_match else 2
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [_escape_markdown_cell(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        width = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (width - len(row)) for row in rows]
        header = normalized_rows[0]
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(["---"] * width) + " |")
        for row in normalized_rows[1:]:
            parts.append("| " + " | ".join(row) + " |")

    return _normalize_extracted_text("\n\n".join(parts))


def _save_extracted_markdown(case_id: int, original_filename: str, text: str) -> str:
    text_dir = os.path.join(UPLOAD_DIR, str(case_id), "doc_text")
    os.makedirs(text_dir, exist_ok=True)
    storage_name = _storage_filename(os.path.splitext(original_filename)[0] + ".md")
    text_path = os.path.join(text_dir, storage_name)
    with open(text_path, "w", encoding="utf-8") as f:
        f.write(text)
    return text_path


def _pdf_page_filename(filename: str) -> str:
    """校验并规范 PDF 转换页文件名。"""
    cleaned = _clean_filename(filename, "")
    if cleaned != filename or not re.match(r"^case\d+_[A-Za-z0-9]+-\d+\.(png|jpg|jpeg)$", cleaned, re.IGNORECASE):
        raise HTTPException(400, "无效的 PDF 页面文件名")
    return cleaned


def _pdf_page_analysis_path(page_path: str) -> str:
    return page_path + ".analysis.json"


def _load_pdf_page_analysis(page_path: str) -> Optional[dict]:
    analysis_path = _pdf_page_analysis_path(page_path)
    if not os.path.exists(analysis_path):
        return None
    try:
        with open(analysis_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _summarize_pdf_page_analysis(analysis: Optional[dict]) -> Optional[dict]:
    """返回给前端的轻量版预分类结果，避免列表接口携带大段 OCR 文本。"""
    if not analysis:
        return None
    summary = dict(analysis)
    prediction = dict(summary.get("prediction") or {})
    if prediction.get("material_type"):
        prediction["material_type_label"] = get_material_type_label(prediction["material_type"])
    if prediction:
        summary["prediction"] = prediction
    ocr_text = summary.pop("ocr_text", "") or ""
    summary["has_ocr_text"] = bool(ocr_text)
    summary["ocr_text_length"] = len(ocr_text)
    return summary


def _save_pdf_page_analysis(page_path: str, analysis: dict) -> None:
    with open(_pdf_page_analysis_path(page_path), "w", encoding="utf-8") as f:
        json.dump(analysis, f, ensure_ascii=False, indent=2)


def _pdf_page_prefix(filename: str) -> str:
    match = re.match(r"^(case\d+_[A-Za-z0-9]+)-\d+\.(?:png|jpg|jpeg)$", filename, re.IGNORECASE)
    return match.group(1) if match else filename


def _prediction_is_weak_default(prediction: Optional[dict]) -> bool:
    if not prediction:
        return True
    if not prediction.get("material_type"):
        return True
    return (
        prediction.get("material_type") == MaterialType.IMAGING_REPORT
        and prediction.get("material_subtype") == "other_imaging_report"
        and float(prediction.get("confidence") or 0) <= 0.75
        and not prediction.get("matched_keywords")
    )


_EXPECTED_MATERIAL_ORDER = {
    MaterialType.ENTRUSTMENT_LETTER: 10,
    MaterialType.ID_CARD: 20,
    MaterialType.TRAFFIC_ACCIDENT_CERT: 30,
    MaterialType.LITIGATION_MATERIAL: 40,
    MaterialType.APPRAISAL_APPLICATION: 50,
    MaterialType.MEDICAL_RECORD: 60,
    MaterialType.IMAGING_REPORT: 70,
}


_ONE_SHOT_MATERIAL_TYPES = {
    MaterialType.ENTRUSTMENT_LETTER,
    MaterialType.ID_CARD,
    MaterialType.TRAFFIC_ACCIDENT_CERT,
    MaterialType.LITIGATION_MATERIAL,
    MaterialType.APPRAISAL_APPLICATION,
}


_OPEN_STAGE_MATERIAL_TYPES = {
    MaterialType.MEDICAL_RECORD,
    MaterialType.IMAGING_REPORT,
}


def _sequence_order(material_type: Optional[str]) -> int:
    return _EXPECTED_MATERIAL_ORDER.get(material_type or "", 999)


def _analysis_plain_text(analysis: Optional[dict]) -> str:
    text = (analysis or {}).get("ocr_text") or ""
    text = text.replace("<nl>", "\n")
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " 图片 ", text)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", "", text)


def _has_standalone_title(text: str, material_type: Optional[str]) -> bool:
    head = text[:360]
    title_map = {
        MaterialType.ENTRUSTMENT_LETTER: ("司法鉴定委托书",),
        MaterialType.ID_CARD: ("居民身份证",),
        MaterialType.TRAFFIC_ACCIDENT_CERT: ("道路交通事故认定书",),
        MaterialType.LITIGATION_MATERIAL: ("民事起诉状",),
        MaterialType.APPRAISAL_APPLICATION: ("鉴定申请书",),
        MaterialType.MEDICAL_RECORD: (
            "住院病案首页", "病案首页", "入院记录", "出院记录", "出院小结",
            "手术记录", "病程记录", "会诊申请单", "会诊记录",
        ),
        MaterialType.IMAGING_REPORT: (
            "检验报告", "检验报告单", "CT", "DR", "CR", "X线", "MRI",
            "磁共振", "超声", "彩超", "肌电图", "影像表现", "诊断意见",
        ),
    }
    return any(keyword in head for keyword in title_map.get(material_type or "", ()))


def _prediction_is_soft(prediction: Optional[dict]) -> bool:
    if _prediction_is_weak_default(prediction):
        return True
    if not prediction:
        return True
    confidence = float(prediction.get("confidence") or 0)
    return confidence < 0.66 or (prediction.get("is_continuation") and not prediction.get("matched_keywords"))


def _should_use_llm_classification(prediction: Optional[dict]) -> bool:
    if not settings.LLM_CLASSIFICATION_ENABLED:
        return False
    if not settings.LLM_API_KEY:
        return False
    if not prediction or not prediction.get("material_type"):
        return True
    return float(prediction.get("confidence") or 0) < settings.LLM_CLASSIFICATION_MIN_CONFIDENCE


def _correct_prediction_by_sequence(prediction: dict, target: dict, reason: str) -> dict:
    """按前后页上下文生成一个校正后的预测，保留当前页 OCR 证据作为提示。"""
    target_type = target.get("material_type")
    corrected = dict(prediction or {})
    corrected.update({
        "material_type": target_type,
        "material_type_label": get_material_type_label(target_type),
        "material_subtype": target.get("material_subtype"),
        "material_subtype_label": target.get("material_subtype_label"),
        "group_name": corrected.get("group_name") or target.get("group_name"),
        "confidence": min(max(float(target.get("confidence") or 0.78), 0.78), 0.88),
        "reason": reason,
        "matched_keywords": corrected.get("matched_keywords") or [],
        "is_continuation": True,
        "sequence_adjusted": True,
    })
    return corrected


def _update_closed_material_types(closed_types: set[str], seen_types: set[str], active_type: Optional[str]) -> None:
    """进入后续阶段后，把前面的一次性材料标记为已结束。

    病历和检查报告属于开放阶段，内部有多个子类和多份报告，不能关闭。
    """
    active_order = _sequence_order(active_type)
    for seen_type in seen_types:
        if seen_type in _ONE_SHOT_MATERIAL_TYPES and _sequence_order(seen_type) < active_order:
            closed_types.add(seen_type)


def _apply_pdf_sequence_context(pages: list[dict]) -> int:
    """利用同一 PDF 的材料顺序和聚堆特征校正页级预分类。

    常见材料顺序是：委托书 -> 身份证 -> 事故认定书 -> 诉讼材料 ->
    鉴定申请书 -> 病历 -> 检查报告。单页关键词容易把病历中的“CT”
    或诉讼材料中的“交通事故认定书”误判成另一类，因此这里用前后页
    和已出现过的阶段做一次保守校正。
    """
    updated = 0
    grouped: dict[str, list[dict]] = {}
    for page in pages:
        grouped.setdefault(_pdf_page_prefix(page["filename"]), []).append(page)

    inherit_types = {
        MaterialType.ENTRUSTMENT_LETTER,
        MaterialType.ID_CARD,
        MaterialType.TRAFFIC_ACCIDENT_CERT,
        MaterialType.LITIGATION_MATERIAL,
        MaterialType.APPRAISAL_APPLICATION,
        MaterialType.MEDICAL_RECORD,
        MaterialType.IMAGING_REPORT,
    }

    for group_pages in grouped.values():
        ordered = sorted(group_pages, key=lambda p: p.get("page_number") or 0)
        page_items = []
        for page in ordered:
            analysis = _load_pdf_page_analysis(page["filepath"])
            if not analysis:
                page_items.append({"page": page, "analysis": None, "prediction": {}, "text": ""})
                continue
            prediction = dict(analysis.get("prediction") or {})
            page_items.append({
                "page": page,
                "analysis": analysis,
                "prediction": prediction,
                "text": _analysis_plain_text(analysis),
            })

        previous_prediction = None
        seen_types: set[str] = set()
        closed_types: set[str] = set()

        for index, item in enumerate(page_items):
            page = item["page"]
            analysis = item["analysis"]
            if not analysis:
                continue
            prediction = dict(item["prediction"])
            current_type = prediction.get("material_type")
            text = item["text"]
            next_prediction = {}
            for next_item in page_items[index + 1:]:
                candidate = next_item["prediction"]
                if candidate.get("material_type") and not _prediction_is_soft(candidate):
                    next_prediction = candidate
                    break

            corrected = None
            if previous_prediction:
                previous_type = previous_prediction.get("material_type")
                next_type = next_prediction.get("material_type")

                if (
                    current_type in closed_types
                    and previous_type
                    and _sequence_order(previous_type) > _sequence_order(current_type)
                ):
                    corrected = _correct_prediction_by_sequence(
                        prediction,
                        previous_prediction,
                        f"{get_material_type_label(current_type)}已在前序材料结束，当前页按所在连续材料校正",
                    )

                elif (
                    previous_type
                    and previous_type == next_type
                    and current_type != previous_type
                    and (
                        _prediction_is_soft(prediction)
                        or _sequence_order(current_type) < _sequence_order(previous_type)
                        or (
                            current_type == MaterialType.IMAGING_REPORT
                            and previous_type == MaterialType.MEDICAL_RECORD
                            and not _has_standalone_title(text, current_type)
                        )
                    )
                ):
                    corrected = _correct_prediction_by_sequence(
                        prediction,
                        previous_prediction,
                        f"根据前后页均为{get_material_type_label(previous_type)}校正",
                    )

                elif (
                    previous_type in inherit_types
                    and prediction.get("is_continuation")
                    and _prediction_is_soft(prediction)
                ):
                    corrected = _correct_prediction_by_sequence(
                        prediction,
                        previous_prediction,
                        "根据前页材料连续性校正：" + (previous_prediction.get("material_type_label") or ""),
                    )

                elif (
                    previous_type == MaterialType.MEDICAL_RECORD
                    and current_type == MaterialType.IMAGING_REPORT
                    and not _has_standalone_title(text, MaterialType.IMAGING_REPORT)
                ):
                    corrected = _correct_prediction_by_sequence(
                        prediction,
                        previous_prediction,
                        "病历连续页中出现检查/影像词，按前页病历材料校正",
                    )

                elif (
                    current_type
                    and current_type in seen_types
                    and _sequence_order(current_type) < _sequence_order(previous_type)
                    and not _has_standalone_title(text, current_type)
                ):
                    corrected = _correct_prediction_by_sequence(
                        prediction,
                        previous_prediction,
                        "根据材料顺序已进入后续类别，按前页连续材料校正",
                    )

            if corrected:
                analysis["prediction"] = corrected
                analysis["classifier_version"] = CLASSIFIER_VERSION
                _save_pdf_page_analysis(page["filepath"], analysis)
                prediction = corrected
                item["prediction"] = corrected
                updated += 1

            if prediction.get("material_type") and not _prediction_is_soft(prediction):
                _update_closed_material_types(closed_types, seen_types, prediction["material_type"])
                previous_prediction = prediction
                seen_types.add(prediction["material_type"])

    return updated


def _load_existing_pdf_page_ocr(page_path: str, save_dir: str) -> tuple[str, Optional[str]]:
    """复用已经保存过的 OCR Markdown，避免删除预分类缓存后重复请求 OCR。"""
    md_path = os.path.join(save_dir, f"{Path(page_path).stem}.md")
    if not os.path.exists(md_path):
        return "", None
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            return f.read(), md_path
    except Exception:
        return "", None


def _ensure_pdf_page_analysis(
    case_id: int,
    page: dict,
    material: Optional[Material],
    save_dir: str,
    db: Session,
    force: bool = False,
) -> tuple[dict, str]:
    """确保 PDF 页已有当前版本预分类结果，返回完整 analysis 和来源。"""
    page_path = page["filepath"]
    cached = _load_pdf_page_analysis(page_path)
    cache_is_fresh = cached and cached.get("classifier_version") == CLASSIFIER_VERSION
    cache_has_ocr_text = bool(cached and (cached.get("ocr_text") or "").strip())

    if cache_is_fresh and cache_has_ocr_text and not force:
        prediction = dict(cached.get("prediction") or {})
        if prediction.get("material_type") in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
            canonical_name, matched_confirmed = _canonicalize_group_name(case_id, prediction.get("group_name"), db)
            if matched_confirmed and canonical_name != prediction.get("group_name"):
                prediction["group_name"] = canonical_name
                prediction["hospital_name_confirmed"] = True
                cached["prediction"] = prediction
                cached["hospital_name_confirmed_at"] = datetime.now().isoformat(timespec="seconds")
                _save_pdf_page_analysis(page_path, cached)
        return cached, "cache"

    try:
        ocr_text = ""
        ocr_file_path = None
        if material and material.ocr_text and material.ocr_status == OcrStatus.COMPLETED:
            ocr_text = material.ocr_text
            ocr_file_path = material.ocr_file_path
        else:
            if cached and cached.get("ocr_text"):
                ocr_text = cached.get("ocr_text") or ""
                ocr_file_path = cached.get("ocr_file_path")
            else:
                ocr_text, ocr_file_path = _load_existing_pdf_page_ocr(page_path, save_dir)
            if not ocr_text:
                ocr_result = run_ocr(page_path, save_dir=save_dir)
                ocr_text = extract_text_from_result(ocr_result)
                ocr_file_path = ocr_result.get("md_path")

        prediction = classify_pdf_page(ocr_text)
        llm_prediction = None
        if ocr_text and _should_use_llm_classification(prediction):
            llm_prediction = classify_pdf_page_with_llm(ocr_text, prediction)
            if llm_prediction and llm_prediction.get("material_type"):
                prediction = llm_prediction
        if prediction.get("material_type") in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
            canonical_name, matched_confirmed = _canonicalize_group_name(case_id, prediction.get("group_name"), db)
            prediction["group_name"] = canonical_name
            if matched_confirmed:
                prediction["hospital_name_confirmed"] = True

        analysis = {
            "ocr_status": "completed" if ocr_text else "failed",
            "ocr_text": ocr_text,
            "ocr_file_path": ocr_file_path,
            "prediction": prediction,
            "classifier_version": CLASSIFIER_VERSION,
            "analyzed_at": datetime.now().isoformat(timespec="seconds"),
        }
        if llm_prediction:
            analysis["llm_classification"] = llm_prediction
        _save_pdf_page_analysis(page_path, analysis)

        if material and ocr_text:
            material.ocr_text = ocr_text
            material.ocr_file_path = ocr_file_path or material.ocr_file_path
            material.ocr_status = OcrStatus.COMPLETED
            material.material_subtype = prediction.get("material_subtype") or classify_material_subtype(material.material_type, ocr_text)

        return analysis, "new" if ocr_text else "failed"
    except Exception as e:
        analysis = {
            "ocr_status": "failed",
            "ocr_text": "",
            "ocr_file_path": None,
            "prediction": classify_pdf_page(""),
            "classifier_version": CLASSIFIER_VERSION,
            "error": str(e),
            "analyzed_at": datetime.now().isoformat(timespec="seconds"),
        }
        _save_pdf_page_analysis(page_path, analysis)
        return analysis, "failed"


_HOSPITAL_SUFFIXES = ("人民医院", "中心医院", "中医院", "医院", "卫生院")


def _normalize_hospital_name_key(name: Optional[str]) -> str:
    text = (name or "").strip()
    text = re.sub(r"[^\u4e00-\u9fa5A-Za-z0-9]", "", text)
    return text.replace("皋", "阜")


def _hospital_name_core(name: Optional[str]) -> str:
    """提取医院名中的中文锚点，尽量去掉 OCR 混入的图形/表格噪声。"""
    key = _normalize_hospital_name_key(name)
    if not key:
        return ""
    matches = re.findall(r"[\u4e00-\u9fa5]{2,30}(?:人民医院|中心医院|中医院|医院|卫生院)", key)
    if matches:
        return max(matches, key=len)
    return key


def _longest_common_substring_len(a: str, b: str) -> int:
    if not a or not b:
        return 0
    previous = [0] * (len(b) + 1)
    best = 0
    for i, char_a in enumerate(a, 1):
        current = [0] * (len(b) + 1)
        for j, char_b in enumerate(b, 1):
            if char_a == char_b:
                current[j] = previous[j - 1] + 1
                best = max(best, current[j])
        previous = current
    return best


def _has_hospital_suffix(name: str) -> bool:
    return any(suffix in (name or "") for suffix in _HOSPITAL_SUFFIXES)


def _hospital_name_similarity(a: str, b: str) -> float:
    ka = _hospital_name_core(a)
    kb = _hospital_name_core(b)
    if not ka or not kb:
        return 0.0
    if ka == kb:
        return 1.0
    score = SequenceMatcher(None, ka, kb).ratio()
    if len(ka) >= 4 and len(kb) >= 4 and (ka in kb or kb in ka):
        score = max(score, 0.82)
    common_len = _longest_common_substring_len(ka, kb)
    common_ratio = common_len / max(1, min(len(ka), len(kb)))
    if common_len >= 4 and (_has_hospital_suffix(ka) or _has_hospital_suffix(kb)):
        score = max(score, min(0.96, 0.68 + common_ratio * 0.28))
    if common_len >= 3 and any(suffix in ka and suffix in kb for suffix in _HOSPITAL_SUFFIXES):
        score = max(score, 0.86)
    return score


def _upsert_hospital_alias(case_id: int, alias_name: Optional[str], canonical_name: Optional[str], db: Session) -> None:
    alias_name = (alias_name or "").strip()
    canonical_name = (canonical_name or "").strip()
    if not alias_name or not canonical_name or alias_name == canonical_name:
        return
    existing = db.query(HospitalNameAlias).filter(
        HospitalNameAlias.case_id == case_id,
        HospitalNameAlias.alias_name == alias_name,
    ).first()
    if existing:
        existing.canonical_name = canonical_name
        return
    db.add(HospitalNameAlias(
        case_id=case_id,
        alias_name=alias_name,
        canonical_name=canonical_name,
    ))


def _confirmed_hospital_names(case_id: int, db: Session) -> list[str]:
    names: list[str] = []
    groups = db.query(MaterialGroup).filter(
        MaterialGroup.case_id == case_id,
        MaterialGroup.material_type.in_([MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]),
        MaterialGroup.is_confirmed == True,  # noqa: E712
    ).all()
    names.extend(g.group_name for g in groups if g.group_name)
    aliases = db.query(HospitalNameAlias).filter(HospitalNameAlias.case_id == case_id).all()
    names.extend(a.canonical_name for a in aliases if a.canonical_name)

    seen = set()
    ordered = []
    for name in names:
        key = _normalize_hospital_name_key(name)
        if key and key not in seen:
            seen.add(key)
            ordered.append(name)
    return ordered


def _canonicalize_group_name(case_id: int, group_name: Optional[str], db: Session) -> tuple[str, bool]:
    """把 OCR/LLM 识别出的医院名收敛到人工确认过的标准医院名。"""
    raw_name = (group_name or "").strip() or "未识别医院"
    if raw_name == "未识别医院":
        return raw_name, False

    alias = db.query(HospitalNameAlias).filter(
        HospitalNameAlias.case_id == case_id,
        HospitalNameAlias.alias_name == raw_name,
    ).first()
    if alias:
        return alias.canonical_name, True

    confirmed_names = _confirmed_hospital_names(case_id, db)
    if not confirmed_names:
        return raw_name, False

    best_name = None
    best_score = 0.0
    for candidate in confirmed_names:
        score = _hospital_name_similarity(raw_name, candidate)
        if score > best_score:
            best_name = candidate
            best_score = score

    if best_name and best_score >= 0.82 and (_has_hospital_suffix(raw_name) or _has_hospital_suffix(best_name)):
        _upsert_hospital_alias(case_id, raw_name, best_name, db)
        return best_name, True

    return raw_name, False


def _replace_aliases_in_text(text: Optional[str], aliases: set[str], canonical_name: str) -> Optional[str]:
    if not text:
        return text
    updated = text
    for alias in sorted((a for a in aliases if a and a != canonical_name), key=len, reverse=True):
        if len(alias) >= 3:
            updated = updated.replace(alias, canonical_name)
    return updated


def _apply_hospital_name_to_material(material: Material, aliases: set[str], canonical_name: str) -> None:
    """把已确认医院名同步到材料 OCR 文本和 PDF 页预分类缓存。"""
    material.ocr_text = _replace_aliases_in_text(material.ocr_text, aliases, canonical_name)
    if material.ocr_file_path and os.path.exists(material.ocr_file_path):
        try:
            with open(material.ocr_file_path, "r", encoding="utf-8") as f:
                content = f.read()
            updated = _replace_aliases_in_text(content, aliases, canonical_name)
            if updated != content:
                with open(material.ocr_file_path, "w", encoding="utf-8") as f:
                    f.write(updated or "")
        except Exception:
            pass

    analysis = _load_pdf_page_analysis(material.file_path)
    if not analysis:
        return
    analysis["ocr_text"] = _replace_aliases_in_text(analysis.get("ocr_text"), aliases, canonical_name) or analysis.get("ocr_text")
    prediction = dict(analysis.get("prediction") or {})
    if prediction:
        if not prediction.get("group_name") or prediction.get("group_name") in aliases:
            prediction["group_name"] = canonical_name
            prediction["hospital_name_confirmed"] = True
        analysis["prediction"] = prediction
    analysis["hospital_name_confirmed_at"] = datetime.now().isoformat(timespec="seconds")
    _save_pdf_page_analysis(material.file_path, analysis)


def _collect_hospital_aliases(case_id: int, canonical_name: str, aliases: set[str], db: Session) -> set[str]:
    """收集同一标准医院名下已经记录过的 OCR 错名，供一次性后台修正。"""
    result = {name.strip() for name in aliases if name and name.strip()}
    if canonical_name:
        result.add(canonical_name)
    stored_aliases = db.query(HospitalNameAlias).filter(
        HospitalNameAlias.case_id == case_id,
        HospitalNameAlias.canonical_name == canonical_name,
    ).all()
    for alias in stored_aliases:
        if alias.alias_name:
            result.add(alias.alias_name)
        if alias.canonical_name:
            result.add(alias.canonical_name)
    return result


def _hospital_name_matches_alias(name: Optional[str], aliases: set[str], canonical_name: str) -> bool:
    name = (name or "").strip()
    if not name:
        return False
    if name in aliases or name == canonical_name:
        return True
    return _hospital_name_similarity(name, canonical_name) >= 0.82


def _sync_group_records_to_hospital(
    *,
    case_id: int,
    target: MaterialGroup,
    aliases: set[str],
    canonical_name: str,
    db: Session,
) -> None:
    """把材料分组、结构化记录、事实库和已生成报告统一到确认后的医院名。"""
    target_materials = db.query(Material).filter(Material.group_id == target.id).all()
    target_material_ids = [material.id for material in target_materials]

    for material in target_materials:
        _apply_hospital_name_to_material(material, aliases, canonical_name)

    if target.material_type == MaterialType.MEDICAL_RECORD:
        record_query = db.query(HospitalRecord).filter(HospitalRecord.case_id == case_id)
        if target_material_ids:
            record_query = record_query.filter(
                (HospitalRecord.group_id == target.id) | (HospitalRecord.material_id.in_(target_material_ids))
            )
        else:
            record_query = record_query.filter(HospitalRecord.group_id == target.id)
        for record in record_query.all():
            record.group_id = target.id
            record.hospital_name = canonical_name

        for event in db.query(MedicalEvent).filter(
            MedicalEvent.case_id == case_id,
            MedicalEvent.group_id == target.id,
        ).all():
            event.hospital_name = canonical_name

    if target.material_type == MaterialType.IMAGING_REPORT:
        report_query = db.query(ImagingReport).filter(ImagingReport.case_id == case_id)
        if target_material_ids:
            report_query = report_query.filter(
                (ImagingReport.group_id == target.id) | (ImagingReport.material_id.in_(target_material_ids))
            )
        else:
            report_query = report_query.filter(ImagingReport.group_id == target.id)
        for report in report_query.all():
            report.group_id = target.id
            report.hospital_name = canonical_name

    for model in (HospitalRecord, ImagingReport, MedicalEvent, UnifiedFact):
        for row in db.query(model).filter(model.case_id == case_id).all():
            if _hospital_name_matches_alias(getattr(row, "hospital_name", None), aliases, canonical_name):
                row.hospital_name = canonical_name

    for fact in db.query(UnifiedFact).filter(UnifiedFact.case_id == case_id).all():
        fact.title = _replace_aliases_in_text(fact.title, aliases, canonical_name)
        fact.summary = _replace_aliases_in_text(fact.summary, aliases, canonical_name)
        fact.source_quote = _replace_aliases_in_text(fact.source_quote, aliases, canonical_name)

    report = db.query(Report).filter(Report.case_id == case_id).first()
    if report:
        for field in ("case_facts", "material_summary", "appraisal_process", "analysis", "opinion"):
            setattr(report, field, _replace_aliases_in_text(getattr(report, field), aliases, canonical_name))


def _renumber_group_materials(group_id: int, db: Session) -> None:
    materials = db.query(Material).filter(Material.group_id == group_id).all()
    materials = sorted(materials, key=material_sequence_key)
    for index, material in enumerate(materials, 1):
        material.page_number = index


def _resolve_group_id(case_id: int, material_type: str, group_id: Optional[int], db: Session) -> Optional[int]:
    if group_id is None:
        return None
    group = db.query(MaterialGroup).filter(
        MaterialGroup.id == group_id,
        MaterialGroup.case_id == case_id,
    ).first()
    if not group:
        raise HTTPException(400, "材料分组不存在")
    if group.material_type == material_type:
        return group.id

    raise HTTPException(400, "材料分组类型不匹配")


def _get_or_create_group_id(case_id: int, material_type: str, group_name: Optional[str], db: Session) -> Optional[int]:
    """病历/检查报告自动导入时按医院创建或复用分组。"""
    if material_type not in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
        return None

    clean_group_name, matched_confirmed = _canonicalize_group_name(case_id, group_name, db)
    existing = db.query(MaterialGroup).filter(
        MaterialGroup.case_id == case_id,
        MaterialGroup.material_type == material_type,
        MaterialGroup.group_name == clean_group_name,
    ).first()
    if existing:
        return existing.id

    max_order = db.query(MaterialGroup).filter(
        MaterialGroup.case_id == case_id,
        MaterialGroup.material_type == material_type,
    ).count()
    group = MaterialGroup(
        case_id=case_id,
        material_type=material_type,
        group_name=clean_group_name,
        sort_order=max_order + 1,
        is_confirmed=matched_confirmed,
    )
    db.add(group)
    db.flush()
    return group.id


def get_material_type_label(t: str) -> str:
    return MaterialType.LABELS.get(t, t)


def material_to_response(m: Material) -> MaterialResponse:
    """统一转换 Material -> MaterialResponse"""
    return MaterialResponse(
        id=m.id,
        case_id=m.case_id,
        material_type=m.material_type,
        material_type_label=get_material_type_label(m.material_type),
        material_subtype=m.material_subtype,
        material_subtype_label=get_material_subtype_label(m.material_subtype),
        group_id=m.group_id,
        description=m.description,
        page_number=m.page_number,
        original_page_number=material_original_page_number(m),
        file_path=m.file_path,
        original_filename=m.original_filename,
        ocr_text=m.ocr_text,
        ocr_status=m.ocr_status,
        ocr_file_path=m.ocr_file_path,
        created_at=m.created_at,
    )


def group_to_response(g: MaterialGroup) -> MaterialGroupResponse:
    """统一转换 MaterialGroup -> MaterialGroupResponse"""
    return MaterialGroupResponse(
        id=g.id,
        case_id=g.case_id,
        material_type=g.material_type,
        group_name=g.group_name,
        sort_order=g.sort_order,
        is_confirmed=bool(g.is_confirmed),
        created_at=g.created_at,
        materials=[material_to_response(m) for m in sorted(g.materials, key=material_sequence_key)],
    )


# ==================== 材料分组 API ====================

@router.post("/groups", response_model=MaterialGroupResponse)
def create_group(data: MaterialGroupCreate, db: Session = Depends(get_db)):
    """创建材料分组（如：新乡市中心医院病历）。"""
    # 自动计算 sort_order
    max_order = db.query(MaterialGroup).filter(
        MaterialGroup.case_id == data.case_id,
        MaterialGroup.material_type == data.material_type
    ).count()

    group = MaterialGroup(
        case_id=data.case_id,
        material_type=data.material_type,
        group_name=data.group_name,
        sort_order=data.sort_order if data.sort_order else max_order + 1,
        is_confirmed=bool(data.is_confirmed),
    )
    db.add(group)
    db.commit()
    db.refresh(group)

    return group_to_response(group)


@router.get("/groups/case/{case_id}", response_model=List[MaterialGroupResponse])
def get_case_groups(case_id: int, material_type: Optional[str] = None, db: Session = Depends(get_db)):
    """获取案件的材料分组"""
    query = db.query(MaterialGroup).filter(MaterialGroup.case_id == case_id)
    if material_type:
        query = query.filter(MaterialGroup.material_type == material_type)
    groups = query.order_by(MaterialGroup.sort_order).all()
    return [group_to_response(g) for g in groups]


@router.put("/groups/{group_id}", response_model=MaterialGroupResponse)
def update_group(group_id: int, data: MaterialGroupUpdate, db: Session = Depends(get_db)):
    """更新分组名称"""
    group = db.query(MaterialGroup).filter(MaterialGroup.id == group_id).first()
    if not group:
        raise HTTPException(404, "分组不存在")
    if data.group_name is not None:
        group.group_name = data.group_name
    if data.sort_order is not None:
        group.sort_order = data.sort_order
    if data.is_confirmed is not None:
        group.is_confirmed = data.is_confirmed
    db.commit()
    db.refresh(group)
    return group_to_response(group)


@router.post("/groups/{group_id}/confirm-hospital-name", response_model=MaterialGroupResponse)
def confirm_hospital_group_name(
    group_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db),
):
    """确认医院标准名称，并把相近 OCR 错名合并到该分组。

    确认后会记录错名 -> 标准名的别名映射，后续自动导入会优先使用标准名。
    """
    group = db.query(MaterialGroup).filter(MaterialGroup.id == group_id).first()
    if not group:
        raise HTTPException(404, "分组不存在")
    if group.material_type not in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
        raise HTTPException(400, "只有病历和检查报告分组需要确认医院名称")

    case_id = group.case_id
    canonical_name = (data.get("canonical_name") or data.get("group_name") or "").strip()
    if not canonical_name:
        raise HTTPException(400, "请输入正确的医院名称")
    merge_similar = bool(data.get("merge_similar", True))

    aliases = {group.group_name, canonical_name}
    target = db.query(MaterialGroup).filter(
        MaterialGroup.case_id == group.case_id,
        MaterialGroup.material_type == group.material_type,
        MaterialGroup.group_name == canonical_name,
    ).first()
    if not target:
        group.group_name = canonical_name
        group.is_confirmed = True
        target = group
    else:
        target.is_confirmed = True

    candidate_groups = [group]
    if merge_similar:
        same_type_groups = db.query(MaterialGroup).filter(
            MaterialGroup.case_id == group.case_id,
            MaterialGroup.material_type == group.material_type,
            MaterialGroup.id != target.id,
        ).all()
        for candidate in same_type_groups:
            if candidate.id == group.id:
                continue
            if not candidate.is_confirmed and _hospital_name_similarity(candidate.group_name, canonical_name) >= 0.82:
                candidate_groups.append(candidate)

    source_group_ids: set[int] = set()
    for source_group in candidate_groups:
        aliases.add(source_group.group_name)
        _upsert_hospital_alias(case_id, source_group.group_name, canonical_name, db)
        source_group_ids.add(source_group.id)
        if source_group.id == target.id:
            continue
        for material in list(source_group.materials):
            material.group_id = target.id
            _apply_hospital_name_to_material(material, aliases, canonical_name)
        db.query(HospitalRecord).filter(HospitalRecord.group_id == source_group.id).update({
            "group_id": target.id,
            "hospital_name": canonical_name,
        })
        if source_group.id != target.id:
            db.delete(source_group)

    target.group_name = canonical_name
    target.is_confirmed = True
    aliases = _collect_hospital_aliases(case_id, canonical_name, aliases, db)
    for source_group_id in source_group_ids:
        if source_group_id == target.id:
            continue
        if target.material_type == MaterialType.MEDICAL_RECORD:
            db.query(MedicalEvent).filter(MedicalEvent.group_id == source_group_id).update({
                "group_id": target.id,
                "hospital_name": canonical_name,
            }, synchronize_session=False)
        if target.material_type == MaterialType.IMAGING_REPORT:
            db.query(ImagingReport).filter(ImagingReport.group_id == source_group_id).update({
                "group_id": target.id,
                "hospital_name": canonical_name,
            }, synchronize_session=False)

    for material in db.query(Material).filter(Material.group_id == target.id).all():
        _apply_hospital_name_to_material(material, aliases, canonical_name)

    db.query(HospitalRecord).filter(
        HospitalRecord.case_id == case_id,
        HospitalRecord.hospital_name.in_(list(aliases)),
    ).update({"hospital_name": canonical_name}, synchronize_session=False)
    db.query(ImagingReport).filter(
        ImagingReport.case_id == case_id,
        ImagingReport.hospital_name.in_(list(aliases)),
    ).update({"hospital_name": canonical_name}, synchronize_session=False)
    db.query(MedicalEvent).filter(
        MedicalEvent.case_id == case_id,
        MedicalEvent.hospital_name.in_(list(aliases)),
    ).update({"hospital_name": canonical_name}, synchronize_session=False)
    db.query(UnifiedFact).filter(
        UnifiedFact.case_id == case_id,
        UnifiedFact.hospital_name.in_(list(aliases)),
    ).update({"hospital_name": canonical_name}, synchronize_session=False)
    _sync_group_records_to_hospital(
        case_id=case_id,
        target=target,
        aliases=aliases,
        canonical_name=canonical_name,
        db=db,
    )
    _renumber_group_materials(target.id, db)
    db.commit()
    db.refresh(target)
    return group_to_response(target)


@router.delete("/groups/{group_id}")
def delete_group(group_id: int, db: Session = Depends(get_db)):
    """删除分组及其下所有材料"""
    group = db.query(MaterialGroup).filter(MaterialGroup.id == group_id).first()
    if not group:
        raise HTTPException(404, "分组不存在")

    # 删除组内所有文件
    for m in group.materials:
        if m.file_path and os.path.exists(m.file_path):
            os.remove(m.file_path)

    db.delete(group)
    db.commit()
    return {"ok": True, "message": f"已删除分组「{group.group_name}」及其中 {len(group.materials)} 个文件"}


# ==================== 材料上传 API ====================

@router.post("/upload/{case_id}", response_model=MaterialResponse)
async def upload_material(
    case_id: int,
    file: UploadFile = File(...),
    material_type: str = Form(...),
    group_id: Optional[int] = Form(None),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    """
    上传单个材料文件

    - material_type: 材料类型
    - group_id: 分组ID（病历/影像需要指定医院分组）
    - description: 文件描述（如：正面、第3页）
    """
    if material_type not in MaterialType.ALL:
        raise HTTPException(400, f"无效的材料类型: {material_type}")
    group_id = _resolve_group_id(case_id, material_type, group_id, db)

    # 保存文件
    case_dir = os.path.join(UPLOAD_DIR, str(case_id))
    os.makedirs(case_dir, exist_ok=True)
    original_filename = _clean_filename(file.filename)
    file_path = os.path.join(case_dir, _storage_filename(original_filename))
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 计算 page_number（同一 group 内排序）
    page_number = None
    if group_id:
        existing_count = db.query(Material).filter(
            Material.group_id == group_id
        ).count()
        page_number = existing_count + 1

    # 自动生成 description（如果未提供）
    if not description and group_id:
        description = f"第{page_number}页" if page_number else None

    material = Material(
        case_id=case_id,
        group_id=group_id,
        material_type=material_type,
        description=description,
        page_number=page_number,
        file_path=file_path,
        original_filename=original_filename,
        ocr_status=OcrStatus.PENDING,
    )
    db.add(material)
    db.commit()
    db.refresh(material)
    return material_to_response(material)


@router.post("/upload-batch/{case_id}", response_model=List[MaterialResponse])
async def upload_materials_batch(
    case_id: int,
    files: List[UploadFile] = File(...),
    material_type: str = Form(...),
    group_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
):
    """
    批量上传材料文件（同一类型、同一分组）
    """
    if material_type not in MaterialType.ALL:
        raise HTTPException(400, f"无效的材料类型: {material_type}")
    group_id = _resolve_group_id(case_id, material_type, group_id, db)

    case_dir = os.path.join(UPLOAD_DIR, str(case_id))
    os.makedirs(case_dir, exist_ok=True)

    # 计算起始页码
    existing_count = 0
    if group_id:
        existing_count = db.query(Material).filter(
            Material.group_id == group_id
        ).count()

    results = []
    for idx, file in enumerate(files):
        original_filename = _clean_filename(file.filename)
        file_path = os.path.join(case_dir, _storage_filename(original_filename))
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        page_number = (existing_count + idx + 1) if group_id else None
        description = f"第{page_number}页" if group_id and page_number else None

        material = Material(
            case_id=case_id,
            group_id=group_id,
            material_type=material_type,
            description=description,
            page_number=page_number,
            file_path=file_path,
            original_filename=original_filename,
            ocr_status=OcrStatus.PENDING,
        )
        db.add(material)
        results.append(material)

    db.commit()
    for m in results:
        db.refresh(m)

    return [material_to_response(m) for m in results]


@router.post("/upload-docx/{case_id}")
async def upload_docx_material(
    case_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """
    上传可编辑 Word 文档，直接解析文字并自动分类导入材料。
    老 .doc 格式不在这里转换，提醒用户另存为 .docx。
    """
    original_filename = _clean_filename(file.filename, "upload.docx")
    lower_name = original_filename.lower()
    if lower_name.endswith(".doc") and not lower_name.endswith(".docx"):
        raise HTTPException(400, "暂不支持 .doc 老格式，请先用 Word/WPS 另存为 .docx 后再导入")
    if not lower_name.endswith(".docx"):
        raise HTTPException(400, "只支持 .docx Word 文档")

    case_dir = os.path.join(UPLOAD_DIR, str(case_id), "documents")
    os.makedirs(case_dir, exist_ok=True)
    docx_path = os.path.join(case_dir, _storage_filename(original_filename))
    with open(docx_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    text = _extract_docx_markdown(docx_path)
    compact_text = re.sub(r"\s+", "", text or "")
    if len(compact_text) < 30:
        raise HTTPException(400, "该 Word 文档未提取到足够文字，可能是图片型文档；请转成 PDF 或图片后导入")

    prediction = classify_pdf_page(text)
    llm_prediction = None
    if _should_use_llm_classification(prediction):
        llm_prediction = classify_pdf_page_with_llm(text, prediction)
        if llm_prediction and llm_prediction.get("material_type"):
            prediction = llm_prediction

    material_type = prediction.get("material_type")
    confidence = float(prediction.get("confidence") or 0)
    if not material_type:
        return {
            "imported": False,
            "filename": original_filename,
            "prediction": prediction,
            "error": "未能判断材料类型，请转为 PDF/图片后使用人工导入，或手动上传到对应材料类别",
        }
    if material_type not in MaterialType.ALL:
        return {
            "imported": False,
            "filename": original_filename,
            "prediction": prediction,
            "error": f"识别出无效材料类型: {material_type}",
        }
    if confidence < 0.5:
        return {
            "imported": False,
            "filename": original_filename,
            "prediction": prediction,
            "error": f"材料类型置信度过低: {confidence:.2f}",
        }

    text_path = _save_extracted_markdown(case_id, original_filename, text)
    group_id = _get_or_create_group_id(
        case_id,
        material_type,
        prediction.get("group_name"),
        db,
    )
    page_number = None
    if group_id:
        page_number = db.query(Material).filter(
            Material.group_id == group_id,
            Material.material_type == material_type,
        ).count() + 1

    material_subtype = prediction.get("material_subtype")
    if not material_subtype:
        material_subtype = classify_material_subtype(material_type, text)

    material = Material(
        case_id=case_id,
        group_id=group_id,
        material_type=material_type,
        material_subtype=material_subtype,
        description=original_filename,
        page_number=page_number,
        file_path=docx_path,
        original_filename=original_filename,
        ocr_text=text,
        ocr_file_path=text_path,
        ocr_status=OcrStatus.COMPLETED,
    )
    db.add(material)
    db.commit()
    db.refresh(material)

    return {
        "imported": True,
        "filename": original_filename,
        "material": material_to_response(material),
        "prediction": {
            **prediction,
            "material_type_label": get_material_type_label(material_type),
            "material_subtype_label": get_material_subtype_label(material_subtype),
        },
        "llm_classification": llm_prediction,
    }


# ==================== 材料查询/删除 API ====================

@router.get("/case/{case_id}", response_model=List[MaterialResponse])
def get_case_materials(case_id: int, material_type: Optional[str] = None, db: Session = Depends(get_db)):
    """获取案件材料列表"""
    query = db.query(Material).filter(Material.case_id == case_id)
    if material_type:
        query = query.filter(Material.material_type == material_type)
    materials = query.order_by(Material.id).all()
    return [material_to_response(m) for m in materials]


@router.get("/case/{case_id}/grouped")
def get_case_materials_grouped(case_id: int, db: Session = Depends(get_db)):
    """
    获取案件材料，按类型分组，分组类型内再按 group 分组
    返回结构：{ type: [ { group: {...}, files: [...] } ] }
    """
    result = {}

    for mt in MaterialType.ALL:
        # 获取该类型所有材料
        all_mats = db.query(Material).filter(
            Material.case_id == case_id,
            Material.material_type == mt
        ).all()

        # 获取该类型的分组
        groups = db.query(MaterialGroup).filter(
            MaterialGroup.case_id == case_id,
            MaterialGroup.material_type == mt
        ).order_by(MaterialGroup.sort_order).all()

        if mt in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
            # 分组类型：按 group 组织
            type_data = []
            group_ids = {g.id for g in groups}
            for g in groups:
                group_mats = [m for m in all_mats if m.group_id == g.id]
                type_data.append({
                    "group": group_to_response(g),
                    "files": [material_to_response(m) for m in sorted(group_mats, key=material_sequence_key)],
                })
            # 没分组或错挂到其他类型分组的孤儿材料，仍要展示，避免 OCR 页漏计
            orphan_mats = [m for m in all_mats if m.group_id is None or m.group_id not in group_ids]
            if orphan_mats:
                type_data.insert(0, {
                    "group": None,
                    "files": [material_to_response(m) for m in orphan_mats],
                })
            result[mt] = type_data
        else:
            # 非分组类型：直接平铺
            result[mt] = [material_to_response(m) for m in all_mats]

    return result


@router.delete("/{material_id}")
def delete_material(material_id: int, db: Session = Depends(get_db)):
    """删除材料文件"""
    material = db.query(Material).filter(Material.id == material_id).first()
    if not material:
        raise HTTPException(404, "材料不存在")

    if material.file_path and os.path.exists(material.file_path):
        os.remove(material.file_path)

    db.delete(material)
    db.commit()
    return {"ok": True}


@router.put("/{material_id}", response_model=MaterialResponse)
def update_material(
    material_id: int,
    description: Optional[str] = None,
    page_number: Optional[int] = None,
    material_subtype: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """更新材料信息"""
    material = db.query(Material).filter(Material.id == material_id).first()
    if not material:
        raise HTTPException(404, "材料不存在")
    if description is not None:
        material.description = description
    if page_number is not None:
        material.page_number = page_number
    if material_subtype is not None:
        material.material_subtype = material_subtype
    db.commit()
    db.refresh(material)
    return material_to_response(material)


@router.put("/{material_id}/move", response_model=MaterialResponse)
def move_material(
    material_id: int,
    data: dict = Body(...),
    db: Session = Depends(get_db),
):
    """将已上传材料转移到其他材料类别或医院分组，用于自动分类后的人工纠错。"""
    material = db.query(Material).filter(Material.id == material_id).first()
    if not material:
        raise HTTPException(404, "材料不存在")

    target_type = data.get("material_type")
    if target_type not in MaterialType.ALL:
        raise HTTPException(400, f"无效的材料类型: {target_type}")

    group_id = data.get("group_id")
    group_name = (data.get("group_name") or "").strip()

    if target_type in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
        if group_name:
            group_id = _get_or_create_group_id(material.case_id, target_type, group_name, db)
        elif group_id:
            group_id = _resolve_group_id(material.case_id, target_type, group_id, db)
        elif material.material_type == target_type and material.group_id:
            group_id = _resolve_group_id(material.case_id, target_type, material.group_id, db)
        else:
            raise HTTPException(400, "病历或检查报告需要选择医院组，或输入新医院名称")
        page_number = db.query(Material).filter(
            Material.group_id == group_id,
            Material.material_type == target_type,
            Material.id != material.id,
        ).count() + 1
    else:
        group_id = None
        page_number = None

    material.material_type = target_type
    material.group_id = group_id
    material.page_number = page_number
    material.material_subtype = classify_material_subtype(target_type, material.ocr_text or "") if material.ocr_text else None
    db.commit()
    db.refresh(material)
    return material_to_response(material)


@router.put("/{material_id}/ocr-text")
def update_ocr_text(material_id: int, data: dict = Body(...), db: Session = Depends(get_db)):
    """更新 OCR 识别文本"""
    material = db.query(Material).filter(Material.id == material_id).first()
    if not material:
        raise HTTPException(404, "材料不存在")
    material.ocr_text = data.get("ocr_text", "")
    material.ocr_status = OcrStatus.COMPLETED
    material.material_subtype = classify_material_subtype(material.material_type, material.ocr_text)
    db.commit()
    return {"ok": True}


@router.post("/case/{case_id}/classify-subtypes")
def classify_case_material_subtypes(case_id: int, db: Session = Depends(get_db)):
    """根据已完成 OCR 的文本重新识别病历/影像子类型。"""
    materials = db.query(Material).filter(
        Material.case_id == case_id,
        Material.material_type.in_([MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]),
        Material.ocr_status == OcrStatus.COMPLETED,
        Material.ocr_text.isnot(None),
        Material.ocr_text != "",
    ).all()

    updated = 0
    results = []
    for material in materials:
        subtype = classify_material_subtype(material.material_type, material.ocr_text)
        if subtype and subtype != material.material_subtype:
            material.material_subtype = subtype
            updated += 1
        results.append({
            "material_id": material.id,
            "filename": material.original_filename,
            "material_type": material.material_type,
            "material_subtype": subtype,
            "material_subtype_label": get_material_subtype_label(subtype),
        })

    db.commit()
    return {"total": len(materials), "updated": updated, "results": results}


# ==================== PDF 转换 API ====================

@router.post("/upload-pdf/{case_id}")
async def upload_and_convert_pdf(
    case_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """
    上传 PDF 文件并转换为 PNG 图片
    - 保存 PDF 文件
    - 调用 pdftoppm 转换为 PNG
    - 返回转换后的页面列表
    """
    from app.utils.pdf_converter import PdfConverter

    pdf_filename = _clean_filename(file.filename, "upload.pdf")
    if not pdf_filename.lower().endswith(".pdf"):
        raise HTTPException(400, "只支持 PDF 文件")

    # 保存 PDF 文件
    case_dir = os.path.join(UPLOAD_DIR, str(case_id))
    os.makedirs(case_dir, exist_ok=True)

    pdf_path = os.path.join(case_dir, _storage_filename(pdf_filename))

    with open(pdf_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 转换为 PNG
    converter = PdfConverter(case_id)
    success, pages, error_msg = converter.convert(pdf_path, original_filename=pdf_filename)

    if not success:
        # 清理上传的 PDF
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        raise HTTPException(500, f"PDF 转换失败: {error_msg}")

    return {
        "pdf_filename": pdf_filename,
        "pdf_path": pdf_path,
        "pages": pages,  # [{page_number, filename, filepath, url}]
        "total_pages": len(pages)
    }


@router.post("/upload-image-pages/{case_id}")
async def upload_image_pages(
    case_id: int,
    files: List[UploadFile] = File(...),
):
    """
    上传一组图片，直接放入 PDF 工作区，复用现有 OCR、预分类和自动导入流程。
    """
    if not files:
        raise HTTPException(400, "请选择要导入的图片")

    try:
        from PIL import Image, ImageOps
    except Exception as exc:
        raise HTTPException(500, f"缺少 Pillow 依赖，无法处理图片: {exc}")

    from app.utils.pdf_converter import PdfConverter

    allowed_exts = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"}
    original_filenames = [_clean_filename(f.filename, f"image_{idx + 1}.png") for idx, f in enumerate(files)]
    for filename in original_filenames:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in allowed_exts:
            raise HTTPException(400, f"不支持的图片格式: {filename}；请使用 png/jpg/jpeg 等常见图片格式")

    converter = PdfConverter(case_id)
    prefix = f"case{case_id}_{uuid.uuid4().hex[:8]}"
    batch_name = original_filenames[0] if len(original_filenames) == 1 else f"图片导入（{len(original_filenames)}张）"

    meta_path = os.path.join(converter.output_dir, prefix + "_meta.json")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump({
            "original_pdf_filename": batch_name,
            "source_type": "image_batch",
            "original_image_filenames": original_filenames,
            "imported_at": datetime.now().isoformat(timespec="seconds"),
        }, f, ensure_ascii=False, indent=2)

    pages = []
    for idx, upload in enumerate(files, start=1):
        content = await upload.read()
        try:
            image = Image.open(BytesIO(content))
            image = ImageOps.exif_transpose(image)
            if image.mode in ("RGBA", "LA") or "transparency" in image.info:
                rgba = image.convert("RGBA")
                background = Image.new("RGB", image.size, (255, 255, 255))
                background.paste(rgba, mask=rgba.getchannel("A"))
                image = background
            else:
                image = image.convert("RGB")
        except Exception as exc:
            raise HTTPException(400, f"图片解析失败: {original_filenames[idx - 1]}，{exc}")

        filename = f"{prefix}-{idx}.png"
        filepath = os.path.join(converter.output_dir, filename)
        image.save(filepath, format="PNG", optimize=True)
        pages.append({
            "page_number": idx,
            "filename": filename,
            "original_pdf_filename": batch_name,
            "source_type": "image_batch",
            "original_image_filename": original_filenames[idx - 1],
            "filepath": filepath,
            "url": f"/uploads/{case_id}/pdf_pages/{filename}",
            "size": os.path.getsize(filepath),
        })

    return {
        "batch_name": batch_name,
        "source_type": "image_batch",
        "pages": pages,
        "total_pages": len(pages),
    }


@router.get("/case/{case_id}/pdf-pages")
def get_pdf_pages(case_id: int, db: Session = Depends(get_db)):
    """
    获取案件的所有 PDF 转换页面（未导入的）
    同时返回已导入的页面信息（来自 Material 表）
    """
    from app.utils.pdf_converter import PdfConverter

    converter = PdfConverter(case_id)
    all_pages = converter.list_pages()

    # 查找已导入到 Material 的页面
    existing_materials = db.query(Material).filter(
        Material.case_id == case_id,
        Material.file_path.like(f"%pdf_pages/%")
    ).all()

    imported_filenames = {m.original_filename: m for m in existing_materials}

    result = []
    for p in all_pages:
        filename = p["filename"]
        material_info = imported_filenames.get(filename)
        analysis = _summarize_pdf_page_analysis(_load_pdf_page_analysis(p["filepath"]))

        if material_info:
            result.append({
                **p,
                "imported": True,
                "material_id": material_info.id,
                "material_type": material_info.material_type,
                "material_type_label": get_material_type_label(material_info.material_type),
                "material_subtype": material_info.material_subtype,
                "material_subtype_label": get_material_subtype_label(material_info.material_subtype),
                "group_id": material_info.group_id,
                "description": material_info.description,
                "analysis": analysis,
            })
        else:
            result.append({
                **p,
                "imported": False,
                "material_id": None,
                "material_type": None,
                "material_type_label": None,
                "material_subtype": None,
                "material_subtype_label": None,
                "group_id": None,
                "description": None,
                "analysis": analysis,
            })

    return result


@router.post("/case/{case_id}/pdf-pages/analyze")
def analyze_pdf_pages(case_id: int, payload: Optional[dict] = Body(None), db: Session = Depends(get_db)):
    """
    对 PDF 工作区页面进行 OCR 和智能预分类。
    默认分析所有未导入页面；传 filenames 可只分析指定页。
    """
    from app.utils.pdf_converter import PdfConverter

    payload = payload or {}
    requested = set(payload.get("filenames") or [])
    force = bool(payload.get("force"))

    converter = PdfConverter(case_id)
    pages = converter.list_pages()
    if requested:
        requested = {_pdf_page_filename(name) for name in requested}
        pages = [p for p in pages if p["filename"] in requested]

    existing_materials = db.query(Material).filter(
        Material.case_id == case_id,
        Material.file_path.like(f"%pdf_pages/%")
    ).all()
    imported_filenames = {m.original_filename: m for m in existing_materials}

    save_dir = os.path.join(UPLOAD_DIR, str(case_id), "ocr_result")
    results = []
    analyzed = 0
    reused = 0
    failed = 0

    for page in pages:
        filename = page["filename"]
        material = imported_filenames.get(filename)
        analysis, source = _ensure_pdf_page_analysis(case_id, page, material, save_dir, db, force=force)

        if source == "cache":
            reused += 1
            results.append({**page, "analysis": _summarize_pdf_page_analysis(analysis), "source": "cache"})
            continue

        if analysis.get("ocr_text"):
            analyzed += 1
        else:
            failed += 1
        results.append({**page, "analysis": _summarize_pdf_page_analysis(analysis), "source": source})

    db.commit()
    sequence_adjusted = _apply_pdf_sequence_context(converter.list_pages())
    for result in results:
        refreshed = _load_pdf_page_analysis(result["filepath"])
        if refreshed:
            result["analysis"] = _summarize_pdf_page_analysis(refreshed)

    return {
        "total": len(pages),
        "analyzed": analyzed,
        "reused": reused,
        "failed": failed,
        "sequence_adjusted": sequence_adjusted,
        "results": results,
    }


@router.post("/case/{case_id}/pdf-pages/auto-import")
def auto_import_pdf_pages_by_prediction(case_id: int, payload: Optional[dict] = Body(None), db: Session = Depends(get_db)):
    """
    OCR 预分类后按预测结果自动导入为材料。
    已导入页面跳过；无法判断类型的页面保留在 PDF 工作区等待人工处理。
    """
    from app.utils.pdf_converter import PdfConverter

    payload = payload or {}
    requested = set(payload.get("filenames") or [])
    force = bool(payload.get("force"))
    min_confidence = float(payload.get("min_confidence") or 0.5)

    converter = PdfConverter(case_id)
    pages = converter.list_pages()
    if requested:
        requested = {_pdf_page_filename(name) for name in requested}
        pages = [p for p in pages if p["filename"] in requested]

    existing_materials = db.query(Material).filter(
        Material.case_id == case_id,
        Material.file_path.like(f"%pdf_pages/%")
    ).all()
    imported_filenames = {m.original_filename: m for m in existing_materials}

    save_dir = os.path.join(UPLOAD_DIR, str(case_id), "ocr_result")
    analyzed = 0
    reused = 0
    ocr_failed = 0
    results = []

    for page in pages:
        filename = page["filename"]
        material = imported_filenames.get(filename)
        if material:
            reused += 1
            results.append({**page, "analysis": _summarize_pdf_page_analysis(_load_pdf_page_analysis(page["filepath"])), "source": "already_imported"})
            continue
        analysis, source = _ensure_pdf_page_analysis(case_id, page, material, save_dir, db, force=force)
        if source == "cache":
            reused += 1
        elif analysis.get("ocr_text"):
            analyzed += 1
        else:
            ocr_failed += 1
        results.append({**page, "analysis": _summarize_pdf_page_analysis(analysis), "source": source})

    db.commit()
    sequence_adjusted = _apply_pdf_sequence_context(converter.list_pages())

    imported = []
    failed = []
    skipped = []

    for page in pages:
        filename = page["filename"]
        page_path = page["filepath"]

        existing = db.query(Material).filter(
            Material.case_id == case_id,
            Material.file_path.like(f"%pdf_pages/{filename}")
        ).first()
        if existing:
            skipped.append({"filename": filename, "reason": "已导入", "material_id": existing.id})
            continue

        analysis = _load_pdf_page_analysis(page_path) or {}
        prediction = analysis.get("prediction") or {}
        material_type = prediction.get("material_type")
        confidence = float(prediction.get("confidence") or 0)
        ocr_text = analysis.get("ocr_text") or ""

        if not material_type:
            failed.append({"filename": filename, "error": "未识别出材料类型"})
            continue
        if material_type not in MaterialType.ALL:
            failed.append({"filename": filename, "error": f"无效材料类型: {material_type}"})
            continue
        if confidence < min_confidence:
            failed.append({"filename": filename, "error": f"置信度过低: {confidence:.2f}"})
            continue

        try:
            group_id = _get_or_create_group_id(
                case_id,
                material_type,
                prediction.get("group_name"),
                db,
            )
            page_number = page.get("page_number")
            if group_id and page_number is None:
                page_number = db.query(Material).filter(
                    Material.group_id == group_id,
                    Material.material_type == material_type,
                ).count() + 1

            predicted_subtype = prediction.get("material_subtype")
            material_subtype = predicted_subtype
            if not material_subtype and ocr_text:
                material_subtype = classify_material_subtype(material_type, ocr_text)

            material = Material(
                case_id=case_id,
                group_id=group_id,
                material_type=material_type,
                description=filename,
                page_number=page_number,
                file_path=page_path,
                original_filename=filename,
                ocr_text=ocr_text or None,
                ocr_file_path=analysis.get("ocr_file_path"),
                ocr_status=OcrStatus.COMPLETED if ocr_text else OcrStatus.PENDING,
                material_subtype=material_subtype,
            )
            db.add(material)
            db.flush()
            db.commit()
            db.refresh(material)
            imported.append({
                "filename": filename,
                "material": material_to_response(material),
                "prediction": _summarize_pdf_page_analysis(analysis).get("prediction"),
            })
        except Exception as e:
            db.rollback()
            failed.append({"filename": filename, "error": f"数据库错误: {str(e)}"})

    db.commit()

    return {
        "total": len(pages),
        "analyzed": analyzed,
        "reused": reused,
        "ocr_failed": ocr_failed,
        "sequence_adjusted": sequence_adjusted,
        "imported": imported,
        "failed": failed,
        "skipped": skipped,
    }


@router.delete("/pdf-page/{case_id}/{filename}")
def delete_pdf_page(case_id: int, filename: str, db: Session = Depends(get_db)):
    """删除 PDF 转换后的单个页面文件"""
    from app.utils.pdf_converter import PdfConverter
    filename = _pdf_page_filename(filename)

    # 检查是否有已导入的 Material
    existing = db.query(Material).filter(
        Material.case_id == case_id,
        Material.file_path.like(f"%pdf_pages/{filename}")
    ).first()

    if existing:
        raise HTTPException(400, "该页面已导入材料，请先撤销导入后再删除")

    converter = PdfConverter(case_id)
    success, msg = converter.delete_page(filename)

    if not success:
        raise HTTPException(404, msg)

    return {"ok": True, "message": msg}


@router.delete("/pdf/{case_id}/{prefix}")
def delete_pdf_all(case_id: int, prefix: str, db: Session = Depends(get_db)):
    """
    删除整个 PDF 的所有转换页面文件
    - prefix: PDF 文件名前缀（如 case123_abc）
    - 同时删除原始 PDF 文件（如果存在）
    """
    from app.utils.pdf_converter import PdfConverter
    if not re.match(r"^case\d+_[A-Za-z0-9]+$", prefix):
        raise HTTPException(400, "无效的 PDF 前缀")

    converter = PdfConverter(case_id)

    # 删除所有转换的页面文件
    deleted_count = 0
    for f in os.listdir(converter.output_dir):
        if f.startswith(prefix) and (f.endswith(".png") or f.endswith(".jpg")):
            filepath = os.path.join(converter.output_dir, f)
            if os.path.isfile(filepath):
                os.remove(filepath)
                deleted_count += 1

    # 删除原始 PDF 文件（如果存在）
    meta_path = os.path.join(converter.output_dir, prefix + "_meta.json")
    pdf_file = os.path.join(UPLOAD_DIR, str(case_id), prefix + ".pdf")
    if os.path.exists(meta_path):
        try:
            with open(meta_path, "r", encoding="utf-8") as mf:
                meta = json.load(mf)
            source_pdf_path = meta.get("source_pdf_path")
            if source_pdf_path:
                pdf_file = source_pdf_path
        except Exception:
            pass
    original_pdf_deleted = False
    if os.path.exists(pdf_file):
        os.remove(pdf_file)
        original_pdf_deleted = True
    if os.path.exists(meta_path):
        os.remove(meta_path)

    return {
        "ok": True,
        "message": f"已删除 PDF {prefix} 的 {deleted_count} 个页面文件" + (" 和原始 PDF 文件" if original_pdf_deleted else ""),
        "deleted_pages": deleted_count,
        "deleted_original_pdf": original_pdf_deleted
    }


@router.post("/import-pdf-pages/{case_id}")
def import_pdf_pages(
    case_id: int,
    filenames: List[str] = Body(...),
    material_type: str = Body(...),
    group_id: int = Body(None),
    group_name: str = Body(None),  # 可选：同时创建新分组
    db: Session = Depends(get_db),
):
    """
    将 PDF 转换页面导入为材料

    - filenames: 要导入的页面文件名列表
    - material_type: 材料类型
    - group_id: 分组ID（病历/影像需要）
    - group_name: 如果没有 group_id，可以传这个来创建新分组
    """
    from app.utils.pdf_converter import PdfConverter

    if material_type not in MaterialType.ALL:
        raise HTTPException(400, f"无效的材料类型: {material_type}")
    group_id = _resolve_group_id(case_id, material_type, group_id, db)

    # 如果需要分组但没有 group_id，先创建分组
    if group_id is None and group_name and material_type in [MaterialType.MEDICAL_RECORD, MaterialType.IMAGING_REPORT]:
        group_name, matched_confirmed = _canonicalize_group_name(case_id, group_name, db)
        # 检查是否已存在同名分组
        existing = db.query(MaterialGroup).filter(
            MaterialGroup.case_id == case_id,
            MaterialGroup.material_type == material_type,
            MaterialGroup.group_name == group_name
        ).first()

        if existing:
            group_id = existing.id
            db.commit()
        else:
            max_order = db.query(MaterialGroup).filter(
                MaterialGroup.case_id == case_id,
                MaterialGroup.material_type == material_type
            ).count()

            new_group = MaterialGroup(
                case_id=case_id,
                material_type=material_type,
                group_name=group_name,
                sort_order=max_order + 1,
                is_confirmed=matched_confirmed,
            )
            db.add(new_group)
            db.commit()
            db.refresh(new_group)
            group_id = new_group.id

    converter = PdfConverter(case_id)
    page_by_filename = {page["filename"]: page for page in converter.list_pages()}
    results = []

    # 计算起始页码
    existing_count = 0
    if group_id:
        existing_count = db.query(Material).filter(
            Material.group_id == group_id,
            Material.material_type == material_type,
        ).count()

    for idx, filename in enumerate(filenames):
        filename = _pdf_page_filename(filename)
        page_path = os.path.join(converter.output_dir, filename)

        if not os.path.exists(page_path):
            results.append({"filename": filename, "success": False, "error": f"文件不存在: {page_path}"})
            continue

        page_info = page_by_filename.get(filename) or {}
        page_number = page_info.get("page_number")
        if group_id and page_number is None:
            page_number = existing_count + idx + 1
        analysis = _load_pdf_page_analysis(page_path) or {}
        prediction = analysis.get("prediction") or {}
        ocr_text = analysis.get("ocr_text") or ""
        ocr_file_path = analysis.get("ocr_file_path")
        predicted_subtype = prediction.get("material_subtype")
        material_subtype = predicted_subtype if predicted_subtype and prediction.get("material_type") == material_type else None
        if not material_subtype and ocr_text:
            material_subtype = classify_material_subtype(material_type, ocr_text)

        # 检查是否已导入
        existing = db.query(Material).filter(
            Material.case_id == case_id,
            Material.file_path.like(f"%pdf_pages/{filename}")
        ).first()

        if existing:
            results.append({"filename": filename, "success": False, "error": "该页面已导入"})
            continue

        try:
            material = Material(
                case_id=case_id,
                group_id=group_id,
                material_type=material_type,
                description=filename,  # 存储文件名作为描述
                page_number=page_number,
                file_path=page_path,
                original_filename=filename,
                ocr_text=ocr_text or None,
                ocr_file_path=ocr_file_path,
                ocr_status=OcrStatus.COMPLETED if ocr_text else OcrStatus.PENDING,
                material_subtype=material_subtype,
            )
            db.add(material)
            db.flush()  # 测试是否这里出问题
            results.append({
                "filename": filename,
                "success": True,
                "material": material_to_response(material)
            })
        except Exception as e:
            db.rollback()
            results.append({"filename": filename, "success": False, "error": f"数据库错误: {str(e)}"})

    db.commit()

    return {
        "imported": [r for r in results if r["success"]],
        "failed": [r for r in results if not r["success"]],
        "group_id": group_id,
    }


@router.post("/revert-import/{case_id}/{filename}")
def revert_pdf_import(case_id: int, filename: str, db: Session = Depends(get_db)):
    """
    撤销 PDF 页面的导入（删除 Material 记录，保留文件）
    """
    filename = _pdf_page_filename(filename)
    material = db.query(Material).filter(
        Material.case_id == case_id,
        Material.file_path.like(f"%pdf_pages/{filename}")
    ).first()

    if not material:
        raise HTTPException(404, "未找到该导入记录")

    material_id = material.id
    db.delete(material)
    db.commit()

    return {"ok": True, "material_id": material_id, "message": "已撤销导入，文件保留"}
