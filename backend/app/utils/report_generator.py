"""
Word 报告生成工具
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path
from typing import Optional

from app.config import settings
from app.models.case import Case, Person, HospitalRecord, ImagingReport, AppSettings
from app.models import Report


APPRAISAL_CENTER_NAME = "河南医药大学司法鉴定中心"
APPRAISAL_CENTER_DOC_PREFIX = "豫河南医药大学司鉴中心"


def generate_report_docx(case: Case, db) -> Path:
    """
    生成司法鉴定意见书 Word 文档
    
    Args:
        case: 案件对象
        db: 数据库会话
    
    Returns:
        生成的文件路径
    """
    doc = Document()
    
    # 设置默认字体（中文）
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 获取关联数据
    person = db.query(Person).filter(Person.case_id == case.id).first()
    report = db.query(Report).filter(Report.case_id == case.id).first()
    hospital_records = db.query(HospitalRecord).filter(HospitalRecord.case_id == case.id).all()
    imaging_reports = db.query(ImagingReport).filter(ImagingReport.case_id == case.id).all()
    
    # ==================== 标题 ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(APPRAISAL_CENTER_NAME)
    run.font.size = Pt(18)
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("司 法 鉴 定 意 见 书")
    run.font.size = Pt(22)
    run.bold = True
    
    # 文号
    doc_num = doc.add_paragraph()
    doc_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_num.add_run(f"{APPRAISAL_CENTER_DOC_PREFIX}{case.case_number or ''}号")
    
    doc.add_paragraph()  # 空行
    
    # ==================== 一、基本情况 ====================
    add_heading(doc, "一、基本情况")
    
    # 1. 委托单位
    add_field(doc, "委托单位", case.entrusting_unit or "待填写")
    # 2. 委托事项
    add_field(doc, "委托事项", case.entrustment_matter or "待填写")
    # 3. 受理日期
    add_field(doc, "受理日期", str(case.acceptance_date) if case.acceptance_date else "")
    # 4. 鉴定材料（材料清单）
    if case.material_list:
        p = doc.add_paragraph()
        p.add_run("鉴定材料：").bold = True
        doc.add_paragraph(case.material_list)
    # 5. 鉴定日期
    add_field(doc, "鉴定日期", str(case.appraisal_date) if case.appraisal_date else "")
    # 6. 鉴定地点
    add_field(doc, "鉴定地点", case.appraisal_location or APPRAISAL_CENTER_NAME)
    # 7. 被鉴定人（合为一行）
    if person:
        name = person.name or "待填写"
        gender = person.gender or ""
        birth = str(person.birth_date) if person.birth_date else ""
        id_num = person.id_number or ""
        addr = person.address or ""
        parts = [name]
        if gender:
            parts.append(gender)
        birth_id = ""
        if birth:
            birth_str = birth.replace("-", "年", 1).replace("-", "月", 1)
            if not birth_str.endswith("日"):
                birth_str += "日"
            birth_id = f"{birth_str}生"
        if id_num:
            birth_id += f"（{id_num}）"
        if birth_id:
            parts.append(birth_id)
        if addr:
            parts.append(f"住{addr}")
        add_field(doc, "被鉴定人", "，".join(parts))
    
    # ==================== 二、基本案情 ====================
    add_heading(doc, "二、基本案情")
    
    if report and report.case_facts:
        add_body_text(doc, report.case_facts)
    elif case.accident_description:
        doc.add_paragraph(case.accident_description)
    else:
        doc.add_paragraph("待根据事故认定书等材料填写...")
    
    # ==================== 三、资料摘要 ====================
    add_heading(doc, "三、资料摘要")
    
    if report and report.material_summary:
        add_body_text(doc, report.material_summary)
    elif hospital_records:
        for i, record in enumerate(hospital_records, 1):
            doc.add_paragraph(format_hospital_record_summary(record, i))
    else:
        doc.add_paragraph("待根据病历材料填写...")
    
    # 不再输出影像学报告到资料摘要
    
    # ==================== 四、鉴定过程 ====================
    add_heading(doc, "四、鉴定过程")
    
    if report and report.appraisal_process:
        add_body_text(doc, report.appraisal_process)
    else:
        doc.add_paragraph("根据《司法鉴定程序通则》及相关鉴定标准，对被鉴定人进行了法医临床检验。")
    
    # ==================== 五、分析说明 ====================
    add_heading(doc, "五、分析说明")
    
    if report and report.analysis:
        add_body_text(doc, report.analysis)
    else:
        doc.add_paragraph("【分析说明由鉴定人根据检验结果和专业标准填写】")
    
    # ==================== 六、鉴定意见 ====================
    add_heading(doc, "六、鉴定意见")
    
    if report and report.opinion:
        add_body_text(doc, report.opinion)
    else:
        doc.add_paragraph("【鉴定意见由鉴定人根据分析说明和专业标准填写】")
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区 - 从配置读取鉴定人信息
    appraiser_name = ""
    appraiser_unit = ""
    for key, attr in [("appraiser_name", "name"), ("appraiser_unit", "unit")]:
        row = db.query(AppSettings).filter(AppSettings.key == key).first()
        if row and row.value:
            if attr == "name":
                appraiser_name = row.value
            elif attr == "unit":
                appraiser_unit = row.value

    if appraiser_name:
        doc.add_paragraph(f"鉴定人：{appraiser_name}")
    else:
        doc.add_paragraph("鉴定人：____________________")
    doc.add_paragraph("复核人：____________________")
    if appraiser_unit:
        doc.add_paragraph(f"鉴定机构：{appraiser_unit}")
    doc.add_paragraph(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
    
    # 保存文件
    settings.REPORT_DIR.mkdir(exist_ok=True)
    filename = f"鉴定意见书_{case.case_number or case.id}.docx"
    file_path = settings.REPORT_DIR / filename
    
    doc.save(file_path)
    return file_path


def add_heading(doc, text: str):
    """添加章节标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_field(doc, label: str, value: str):
    """添加字段"""
    p = doc.add_paragraph()
    p.add_run(f"{label}：").bold = True
    p.add_run(value)


def add_body_text(doc, text: str):
    """按自然段写入正文，避免整段文本挤在一个段落里。"""
    text = (text or "").strip()
    if not text:
        return
    for block in text.split("\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)


def _sentence(text: Optional[str]) -> str:
    text = (text or "").strip(" \n\t。；;,，")
    if not text:
        return ""
    return text if text.endswith(("。", "；", "！", "？")) else f"{text}。"


def format_hospital_record_summary(record: HospitalRecord, index: int) -> str:
    """Word 兜底用的病历摘要；正式内容优先使用 report.material_summary。"""
    cn_nums = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）"]
    prefix = cn_nums[index - 1] if index <= len(cn_nums) else f"（{index}）"
    hospital = record.hospital_name or "某医院"
    admission_number = f"（住院号：{record.admission_number}）" if record.admission_number else ""
    text = f"{prefix}据{hospital}住院病案{admission_number}记载："

    date_bits = []
    if record.admission_date:
        date_bits.append(f"{record.admission_date}入院")
    if record.discharge_date:
        date_bits.append(f"{record.discharge_date}出院")
    if record.hospital_days:
        date_bits.append(f"住院{record.hospital_days}天")
    if date_bits:
        text += "，".join(date_bits) + "。"

    if record.chief_complaint:
        text += f"入院时以“{record.chief_complaint}”为主诉。"
    if record.admission_diagnosis:
        text += f"入院诊断为：{record.admission_diagnosis}。"
    if record.treatment_process:
        text += _sentence(f"住院期间予以{record.treatment_process}")
    if record.discharge_diagnosis:
        text += f"出院诊断为：{record.discharge_diagnosis}。"
    if record.discharge_orders:
        text += f"出院医嘱：{record.discharge_orders}。"
    return text
